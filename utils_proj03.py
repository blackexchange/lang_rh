
import os
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate, HumanMessagePromptTemplate, MessagesPlaceholder
import json
import pandas as pd
import csv
import streamlit as st
import re
from io import BytesIO

# Importa PyMuPDF (mais simples e confiável)
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# Tenta importar docling como opção alternativa (pode ter problemas de permissão no Windows)
try:
    from docling.document_converter import DocumentConverter
    DOCLING_AVAILABLE = True
except Exception:
    DOCLING_AVAILABLE = False

# Verifica se pelo menos uma biblioteca está disponível
if not PYMUPDF_AVAILABLE and not DOCLING_AVAILABLE:
    import sys
    print("ERRO: Nenhuma biblioteca de PDF disponível. Instale PyMuPDF: pip install PyMuPDF")
    sys.exit(1)

def load_llm(id_model, temperature):
  llm = ChatGroq(
      model=id_model,
      temperature=temperature,
      max_tokens=None,
      timeout=None,
      max_retries=2,
  )
  return llm


def format_res(res, return_thinking=False):
  res = res.strip()

  if return_thinking:
    res = res.replace("<think>", "[pensando...] ")
    res = res.replace("</think>", "\n---\n")

  else:
    if "</think>" in res:
      res = res.split("</think>")[-1].strip()

  return res


def parse_doc(file_path):
  """
  Extrai texto de um arquivo PDF usando PyMuPDF (padrão) ou docling como alternativa.
  
  Args:
    file_path: Caminho para o arquivo PDF
    
  Returns:
    str: Conteúdo do PDF em formato texto
  """
  # Usa PyMuPDF por padrão (mais simples e confiável)
  if PYMUPDF_AVAILABLE:
    try:
      doc = fitz.open(file_path)
      content = ""
      for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        if text:
          content += f"\n--- Página {page_num + 1} ---\n"
          content += text
      doc.close()
      return content.strip()
    except Exception as e:
      st.error(f"Erro ao processar PDF com PyMuPDF: {e}")
      # Tenta docling como fallback se PyMuPDF falhar
      if DOCLING_AVAILABLE:
        st.info("Tentando docling como alternativa...")
      else:
        raise
  
  # Fallback para docling (se PyMuPDF não estiver disponível ou falhar)
  if DOCLING_AVAILABLE:
    try:
      # Configura variável de ambiente para evitar problemas de cache
      os.environ.setdefault('HF_HOME', os.path.join(os.getcwd(), '.hf_cache'))
      
      converter = DocumentConverter()
      result = converter.convert(file_path)
      content = result.document.export_to_markdown()
      return content
    except (OSError, PermissionError) as e:
      st.error(f"Erro de permissão com docling: {e}")
      st.info("Sugestão: Use PyMuPDF que não requer downloads de modelos.")
      raise
    except Exception as e:
      st.error(f"Erro ao processar PDF com docling: {e}")
      raise
  
  # Se nenhuma biblioteca estiver disponível
  raise ImportError("Nenhuma biblioteca de PDF disponível. Instale PyMuPDF: pip install PyMuPDF")


def parse_res_llm(response_text: str, required_fields: list) -> dict:
    try:
        # Remove a parte do raciocínio (<think>...</think>)
        if "</think>" in response_text:
            response_text = response_text.split("</think>")[-1].strip()

        # Localiza o JSON e faz o parsing
        start_idx = response_text.find('{')
        end_idx = response_text.rfind('}') + 1
        if start_idx == -1 or end_idx == 0:
            raise json.JSONDecodeError("Nenhum JSON encontrado na resposta", response_text, 0)

        json_str = response_text[start_idx:end_idx]
        info_cv = json.loads(json_str)

        for field in required_fields:
            if field not in info_cv:
                info_cv[field] = []

        return info_cv

    except json.JSONDecodeError:
        #Erro ao interpretar a resposta do modelo
        return


def save_json_cv(new_data, path_json, key_name="name"):
    # Carrega o JSON existente, se houver
    if os.path.exists(path_json):
        with open(path_json, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = []

    if isinstance(data, dict):
        data = [data]

    # Verifica se já existe um currículo com o mesmo nome
    candidates = [entry.get(key_name) for entry in data]
    if new_data.get(key_name) in candidates:
        st.warning(f"Currículo '{new_data.get(key_name)}' já registrado. Ignorando.")
        return

    # Adiciona e salva
    data.append(new_data)
    with open(path_json, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def load_json_cv(path_json):
    with open(path_json, "r", encoding="utf-8") as f:
        return json.load(f)


def generate_cv_content_from_json(cv_data):
    """
    Gera conteúdo de currículo em texto a partir dos dados estruturados do JSON.
    
    Args:
        cv_data: Dicionário com os dados do currículo do JSON
    
    Returns:
        str: Conteúdo do currículo em formato texto
    """
    content = f"""# {cv_data.get('name', 'Currículo')}

    ## Resumo Profissional
    {cv_data.get('summary', '')}

    ## Experiências
    {cv_data.get('experiences', '')}

    ## Formação Acadêmica
    {cv_data.get('academic_info', '')}

    ## Certificações
    {cv_data.get('certifications', '')}

    ## Cursos
    {cv_data.get('training_courses', '')}
    
    ## Hard Skills
    {cv_data.get('hard_skills', '')}

    ## Soft Skills
    {cv_data.get('soft_skills', '')}

    """
    return content


def generate_analysis_from_json(cv_data):
    """
    Gera uma análise estruturada a partir dos dados do JSON.
    
    Args:
        cv_data: Dicionário com os dados do currículo do JSON
    
    Returns:
        dict: Análise estruturada
    """
    analysis = {
        "analysis_summary": f"Análise do perfil de {cv_data.get('name', 'candidato')}. {cv_data.get('summary', '')}",
        "alignment_score": float(cv_data.get('score', 0.0)),
        "strengths": cv_data.get('strengths', []),
        "weaknesses": cv_data.get('areas_for_development', []),
        "missing_skills": [],
        "underutilized_skills": [],
        "recommendations": [cv_data.get('final_recommendations', '')],
        "key_improvements": cv_data.get('important_considerations', [])
    }
    return analysis


def show_cv_result(result: dict):
    md = f"### 📄 Análise e Resumo do Currículo\n"
    if "name" in result:
        md += f"- **Nome:** {result['name']}\n"
    if "area" in result:
        md += f"- **Área de Atuação:** {result['area']}\n"
    if "skills" in result:
        md += f"- **Competências:** {', '.join(result['skills'])}\n"
    if "summary" in result:
        md += f"- **Resumo do Perfil:** {result['summary']}\n"
    if "interview_questions" in result:
        md += f"- **Perguntas sugeridas:**\n"
        md += "\n".join([f"  - {q}" for q in result["interview_questions"]]) + "\n"
    if "strengths" in result:
        md += f"- **Pontos fortes (ou Alinhamentos):**\n"
        md += "\n".join([f"  - {s}" for s in result["strengths"]]) + "\n"
    if "areas_for_development" in result:
        md += f"- **Pontos a desenvolver (ou Desalinhamentos):**\n"
        md += "\n".join([f"  - {a}" for a in result["areas_for_development"]]) + "\n"
    if "important_considerations" in result:
        md += f"- **Pontos de atenção:**\n"
        md += "\n".join([f"  - {i}" for i in result["important_considerations"]]) + "\n"
    if "final_recommendations" in result:
        md += f"- **Conclusão e recomendações:** {result['final_recommendations']}\n"
    return md


def save_job_to_csv(data, filename):
    headers = ['title', 'description', 'details']
    file_exists = os.path.exists(filename)

    with open(filename, 'a', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=headers, delimiter=';')
        if not file_exists:
            writer.writeheader()
        writer.writerow(data)


def load_job(csv_path):
  try:
    df = pd.read_csv(csv_path, sep=';', encoding='utf-8')
    job = df.iloc[-1]
    prompt_text = f"""
    **Vaga para {job['title']}**

    **Descrição da Vaga:**
    {job['description']}

    **Detalhes Completos:**
    {job['details']}
    """

    return prompt_text.strip()

  except FileNotFoundError:
    return "Erro: Arquivo de vagas não encontrado"

def process_cv(schema, job_details, prompt_template, prompt_score, llm, file_path):

  if file_path:
    if not os.path.exists(file_path):
      raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")

  content = parse_doc(file_path)

  chain = prompt_template | llm
  output = chain.invoke({"schema": schema, "cv": content, "job": job_details, "prompt_score": prompt_score})

  res = format_res(output.content)

  return output, res


def display_json_table(path_json):
  with open(path_json, "r", encoding="utf-8") as f:
    data = json.load(f)

  df = pd.DataFrame(data)
  return df


# ============================================
# AGENTE ANALISADOR - Analisa currículo e vaga
# ============================================

def create_analysis_prompt_template():
    """Cria o template de prompt para o agente analisador"""
    return ChatPromptTemplate.from_template("""
    Você é um especialista em Recursos Humanos com vasta experiência em análise de currículos.
    Sua tarefa é analisar profundamente o currículo e a vaga, gerando uma análise detalhada e estruturada.

    INSTRUÇÕES:
    1. Analise o currículo fornecido em detalhes
    2. Compare com os requisitos e características da vaga
    3. Identifique pontos fortes, fracos, alinhamentos e desalinhamentos
    4. Gere recomendações específicas para melhorar o currículo
    5. Retorne APENAS um JSON válido com a estrutura abaixo

    SCHEMA DE RESPOSTA (JSON):
    {{
    "analysis_summary": "Resumo executivo da análise (2-3 parágrafos)",
    "alignment_score": 0.0,
    "strengths": [
        "Lista de pontos fortes e alinhamentos com a vaga"
    ],
    "weaknesses": [
        "Lista de pontos fracos e desalinhamentos com a vaga"
    ],
    "missing_skills": [
        "Habilidades mencionadas na vaga que não estão no currículo"
    ],
    "underutilized_skills": [
        "Habilidades do candidato que poderiam ser melhor destacadas"
    ],
    "recommendations": [
        "Recomendações específicas para melhorar o currículo"
    ],
    "key_improvements": [
        "Melhorias prioritárias que devem ser feitas no currículo"
    ]
    }}

    CURRÍCULO:
    '{cv}'

    VAGA:
    '{job}'

    Retorne APENAS o JSON, sem explicações adicionais.
    """)


def analyze_cv_and_job(llm, cv_content, job_details):
    """
    Agente Analisador: Analisa o currículo e a vaga, gerando análise detalhada
    
    Args:
        llm: Modelo de linguagem
        cv_content: Conteúdo do currículo (texto/markdown)
        job_details: Detalhes da vaga (texto)
    
    Returns:
        dict: Análise estruturada em JSON
    """
    prompt_template = create_analysis_prompt_template()
    chain = prompt_template | llm
    
    output = chain.invoke({
        "cv": cv_content,
        "job": job_details
    })
    
    res = format_res(output.content)
    
    # Parse do JSON
    try:
        start_idx = res.find('{')
        end_idx = res.rfind('}') + 1
        if start_idx == -1 or end_idx == 0:
            raise json.JSONDecodeError("Nenhum JSON encontrado", res, 0)
        
        json_str = res[start_idx:end_idx]
        analysis = json.loads(json_str)
        return analysis
    except json.JSONDecodeError as e:
        st.error(f"Erro ao processar análise: {e}")
        return None


# ============================================
# AGENTE REFORMULADOR - Reformula o currículo
# ============================================

def create_rewrite_prompt_template():
    """Cria o template de prompt para o agente reformulador"""
    return ChatPromptTemplate.from_template("""
    Atue como um Especialista em Recrutamento e Seleção com 15 anos de experiência em recolocação profissional. Sua tarefa é reescrever o  e Otimização de Currículos para ATS. Você é expert em adequar perfis seniores (Arquitetos, Tech Leads, Engenheiros de Dados, Especialistas em IA Generativa) para vagas competitivas.
    Você irá reescrever pontos estratégicos do currículo abaixo usando o TEMPLATE fornecido como base estrutural, para maximizar a aderência (match) com a descrição da vaga fornecida, baseadas na *análise* e no currículo original.. Pode inventar dados, alterar a semântica e a ênfase para espelhar a linguagem da empresa.


    INSTRUÇÕES IMPORTANTES:
    1. Use o TEMPLATE fornecido como estrutura base - mantenha a mesma formatação, seções e estilo
    2. Preencha as seções do template com informações baseadas na análise e no **currículo original**
    3. Você PODE e DEVE criar/inventar conteúdo relevante para as seções, desde que seja coerente com:
    - As habilidades e experiências do candidato
    - Os pontos fortes e fracos identificados na análise
    - Os requisitos da vaga
    - As recomendações da análise
    - As habilidades faltantes identificadas na análise
    - As habilidades subutilizadas identificadas na análise
    4. Seções a preencher:
    - **Resumo Profissional**: 
        "Resumo Profissional" de um currículo para torná-lo altamente magnético e otimizado para uma vaga específica.
        Instruções:
        - Analise a Descrição da Vaga: Identifique as 3 competências técnicas (hard skills) e as 2 competências comportamentais (soft skills) mais críticas.
        - Analise o Currículo: Extraia as conquistas e experiências que mais se conectam com as necessidades da vaga.
        - Redação do Resumo: Escreva um parágrafo de 4 a 6 linhas que siga a estrutura:
            - [Título Profissional] + [Tempo de Experiência] + [Principal Especialidade].
            - Mencione uma conquista quantificável (números, % ou resultados).
            - Inclua as palavras-chave principais da vaga de forma natural.
            Diretrizes de Estilo:
                - Use tom profissional, direto e focado em resultados.
                - Evite clichês como "apaixonado por desafios" ou "proativo".
                - Use verbos de ação (Ex: Gerenciei, Desenvolvi, Reduzi, Implementei).

    - **Experiências**: Selecione as experiências profissionais do currículo que sejam mais relevantes para a vaga. Não invente o nome das empresas nem a data de início e fim das experiências. Adeque o nome da posição/cargo ao da vaga
        - Adeque cada experiência profissional do currículo, para ficarem aderente aos requisitos da vaga.
        - Para cada uma, melhore, crie 2 casos mais relevantes, contextualizados com a respectiva área.
        - Reescreva 3 pontos para cada uma, focados no Modelo STAR: "Contexto + Verbo de Ação + Contexto Técnico + Resultado".
        - As posições/cargos das experiências profissionais do currículo devem ser adequadas ao da vaga.
    - **Formação Acadêmica**:
        - Não invente a instituição de ensino nem o título do curso. Apenas liste as formações acadêmicas do currículo.
    - **Projetos e Consultorias Relevantes**: Crie 2 a 4 projetos, contendo nome de empresa e data fake, que demonstrem as habilidades necessárias para a vaga. O nome da empresa deve ser criativo (evite X, YZ etc...).
    - **Hard Skills**: Identifique as hard skills da vaga e adicione-os, inserindo mais outras adjacentes para completar as 4 a 6 hard skills.
    - **Soft Skills**: Identifique as soft skills da vaga e adicione-os, inserindo mais outras adjacentes para completar as 4 a 6 soft skills.
    - **Certificações**: Liste de 4 a 6 certificações do curriculum que estejam alinhadas direta ou indiretamente com a vaga. não invente certificações
    - **Cursos e Treinamentos**: Liste de 4 a 6 cursos e treinamentos do curriculum relevantes, não invente cursos e treinamentos.
    5. Use linguagem {style}
    6. Mantenha a estrutura e formatação exata do template
    7. {focus_instruction}
    8. {highlight_instruction}
    9. {strengths_instruction}
    10. Seja criativo mas realista - crie conteúdo que faça sentido para o perfil do candidato
    11. Traduza para o idioma {idioma}

    TEMPLATE DE CV (use esta estrutura):
    '{cv_template}'

    CURRÍCULO ORIGINAL (referência de informações):
    '{original_cv}'

    ANÁLISE REALIZADA:
    '{analysis}'

    VAGA DE REFERÊNCIA:
    '{job}'

    Crie um currículo completo preenchendo o template com informações relevantes baseadas na análise e no currículo original.
    Mantenha a estrutura exata do template, apenas preenchendo as seções com conteúdo novo e relevante.
    Retorne o currículo completo no mesmo formato do template.
    
    **Omita quaisquer observações ou comentários.**
    **Não invente o nome das empresas nem a data de início e fim das experiências.**
    """)


def rewrite_cv(llm, original_cv_content, analysis, job_details, cv_template=None, rewrite_options=None, idioma="Português Brasileiro"):
    """
    Agente Reformulador: Reformula o currículo baseado na análise usando um template
    
    Args:
        llm: Modelo de linguagem
        original_cv_content: Conteúdo original do currículo
        analysis: Análise gerada pelo agente analisador (dict)
        job_details: Detalhes da vaga
        cv_template: Template de CV para usar como estrutura base (opcional)
        rewrite_options: Dicionário com opções de reformulação (opcional)
        idioma: Idioma do currículo (opcional)
    Returns:
        str: Currículo reformulado em markdown
    """
    # Validação de entrada
    if not original_cv_content:
        raise ValueError("Conteúdo do currículo original não pode estar vazio")
    if not analysis:
        raise ValueError("Análise não pode estar vazia")
    if not job_details:
        raise ValueError("Detalhes da vaga não podem estar vazios")
    if not cv_template:
        raise ValueError("Template de CV é obrigatório")
    
    # Opções padrão
    if rewrite_options is None:
        rewrite_options = {
            "focus": "all",
            "style": "professional",
            "highlight_missing": True,
            "emphasize_strengths": True,
            "template": "1"
        }
    
    # Define instruções baseadas nas opções
    focus_map = {
        "all": "Foque em todos os aspectos do currículo",
        "skills": "Dê ênfase especial às habilidades e competências técnicas",
        "experience": "Dê ênfase especial à experiência profissional e histórico de trabalho",
        "summary": "Dê ênfase especial ao resumo profissional e objetivo"
    }
    focus_instruction = focus_map.get(rewrite_options.get("focus", "all"), focus_map["all"])
    
    highlight_instruction = ""
    if rewrite_options.get("highlight_missing", True):
        highlight_instruction = "Destaque claramente as habilidades mencionadas na vaga que estão faltando no currículo, mas NÃO invente que o candidato as possui."
    else:
        highlight_instruction = "Não é necessário destacar habilidades faltantes."
    
    strengths_instruction = ""
    if rewrite_options.get("emphasize_strengths", True):
        strengths_instruction = "Enfatize e destaque os pontos fortes e alinhamentos identificados na análise."
    else:
        strengths_instruction = "Mantenha os pontos fortes sem ênfase especial."
    
    # Converte a análise para texto estruturado
    analysis_text = f"""
    RESUMO DA ANÁLISE:
    {analysis.get('analysis_summary', 'N/A')}

    PONTOS FORTES:
    {chr(10).join(['- ' + str(s) for s in analysis.get('strengths', [])])}

    PONTOS FRACOS:
    {chr(10).join(['- ' + str(w) for w in analysis.get('weaknesses', [])])}

    HABILIDADES FALTANTES:
    {chr(10).join(['- ' + str(s) for s in analysis.get('missing_skills', [])])}

    HABILIDADES SUBUTILIZADAS:
    {chr(10).join(['- ' + str(s) for s in analysis.get('underutilized_skills', [])])}

    RECOMENDAÇÕES:
    {chr(10).join(['- ' + str(r) for r in analysis.get('recommendations', [])])}

    MELHORIAS PRIORITÁRIAS:
    {chr(10).join(['- ' + str(k) for k in analysis.get('key_improvements', [])])}
    """
    
    try:
        prompt_template = create_rewrite_prompt_template()
        chain = prompt_template | llm
        
        style_map = {
            "professional": "profissional e objetiva",
            "modern": "moderna e dinâmica",
            "concise": "concisa e direta"
        }
        style_text = style_map.get(rewrite_options.get("style", "professional"), "profissional e objetiva")
        
        output = chain.invoke({
            "cv_template": cv_template,
            "original_cv": original_cv_content,
            "analysis": analysis_text,
            "job": job_details,
            "style": style_text,
            "focus_instruction": focus_instruction,
            "highlight_instruction": highlight_instruction,
            "strengths_instruction": strengths_instruction,
            "idioma": idioma
        })
        
        rewritten_cv = format_res(output.content)
        
        # Validação de saída
        if not rewritten_cv:
            raise ValueError("Resposta do modelo está vazia")
        
        rewritten_cv = rewritten_cv.strip()
        
        if len(rewritten_cv) < 50:
            raise ValueError(f"Currículo reformulado está muito curto ({len(rewritten_cv)} caracteres). Mínimo esperado: 50 caracteres")
        
        return rewritten_cv
    except Exception as e:
        # Não usa st.error aqui pois pode não estar no contexto do Streamlit
        error_msg = f"Erro ao reformular currículo: {str(e)}"
        print(error_msg)  # Log para debug
        raise Exception(error_msg) from e


def save_rewritten_cv(content, filename):
    """Salva o currículo reformulado em um arquivo markdown"""
    with open(filename, "w", encoding="utf-8") as f:
        f.write(content)
    return filename


def convert_markdown_to_text(markdown_text):
    """
    Converte markdown básico para texto simples para PDF.
    Remove formatação markdown e mantém estrutura.
    """
    # Remove markdown bold (**texto**)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', markdown_text)
    # Remove markdown italic (*texto*)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Remove markdown headers (# ## ###)
    text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
    # Remove markdown links [text](url)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    # Remove markdown horizontal rules (---)
    text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
    
    return text


def generate_pdf_from_cv(cv_content, filename=None, primary_color="#2563eb"):
    """
    Gera um PDF profissional a partir do conteúdo do currículo em markdown.
    
    Características:
    - 4 níveis de hierarquia tipográfica
    - Ícones para cada seção
    - Cabeçalho colorido com dados de contato
    - Cores temáticas personalizáveis
    - Quebra de página inteligente (evita seções órfãs)
    - Espaçamento 1.5
    - Fonte tamanho 11 (normal)
    - Fonte moderna (Helvetica)
    
    Args:
        cv_content: Conteúdo do currículo em markdown/texto
        filename: Nome do arquivo (opcional)
        primary_color: Cor predominante em hex (ex: "#2563eb")
    
    Returns:
        bytes: Conteúdo do PDF em bytes
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib import colors
        
        # Mapeamento de seções para ícones
        SECTION_ICONS = {
            "Resumo Profissional": "",
            "Resumo": "",
            "Experiências": "",
            "Experiência": "",
            "Formação Acadêmica": "",
            "Formação": "",
            "Certificações": "",
            "Certificação": "",
            "Cursos": "",
            "Cursos e Treinamentos": "",
            "Projetos e Consultorias Relevantes": "",
            "Projetos": "",
            "Hard Skills": "",
            "Habilidades Técnicas": "",
            "Soft Skills": "",
            "Habilidades Comportamentais": ""
        }
        
        def get_section_icon(section_name):
            """Retorna o ícone para uma seção"""
            for key, icon in SECTION_ICONS.items():
                if key.lower() in section_name.lower():
                    return icon
            return "📋"  # Ícone padrão
        
        def hex_to_color(hex_color):
            """Converte cor hex para objeto colors do reportlab"""
            hex_color = hex_color.lstrip('#')
            return colors.HexColor(f"#{hex_color}")
        
        def escape_xml(text):
            """Escapa caracteres especiais para XML, preservando tags HTML"""
            # Protege tags HTML existentes
            text = text.replace('<b>', '___BOLD_START___')
            text = text.replace('</b>', '___BOLD_END___')
            text = text.replace('<i>', '___ITALIC_START___')
            text = text.replace('</i>', '___ITALIC_END___')
            text = text.replace('<br/>', '___BR___')
            
            # Escapa caracteres especiais
            text = text.replace('&', '&amp;')
            text = text.replace('<', '&lt;')
            text = text.replace('>', '&gt;')
            
            # Restaura tags HTML
            text = text.replace('___BOLD_START___', '<b>')
            text = text.replace('___BOLD_END___', '</b>')
            text = text.replace('___ITALIC_START___', '<i>')
            text = text.replace('___ITALIC_END___', '</i>')
            text = text.replace('___BR___', '<br/>')
            
            return text
        
        def process_markdown(text):
            """Processa markdown básico para HTML"""
            # Processa negrito **texto**
            text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
            # Processa itálico *texto* (mas não se já está em negrito)
            text = re.sub(r'(?<!<b>)\*([^*<]+?)\*(?!</b>)', r'<i>\1</i>', text)
            # Processa links [text](url) -> text
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            return text
        
        # Cria buffer em memória
        buffer = BytesIO()
        
        # Cria documento PDF com margens
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=3*cm,  # Espaço para cabeçalho
            bottomMargin=2*cm
        )
        
        # Estilos
        styles = getSampleStyleSheet()
        base_font = 'Helvetica'
        base_font_bold = 'Helvetica-Bold'
        base_font_italic = 'Helvetica-Oblique'
        base_font_bold_italic = 'Helvetica-BoldOblique'
        
        # Espaçamento 1.5 (leading = fontSize * 1.5)
        base_font_size = 11
        base_leading = base_font_size * 1.5
        
        # Nível 1: Nome do profissional (maior, negrito)
        name_style = ParagraphStyle(
            'NameStyle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=12,
            spaceBefore=0,
            alignment=TA_LEFT,
            fontName=base_font_bold,
            leading=36
        )
        
        # Nível 2: Título de seção (com ícone, sem background)
        section_style = ParagraphStyle(
            'SectionStyle',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=hex_to_color(primary_color),  # Usa a cor como texto, não background
            spaceAfter=8,
            spaceBefore=16,
            alignment=TA_LEFT,
            fontName=base_font_bold,
            leading=21
        )
        
        # Nível 3: Subtítulo de seção (experiências, projetos)
        subtitle_style = ParagraphStyle(
            'SubtitleStyle',
            parent=styles['Heading3'],
            fontSize=13,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=6,
            spaceBefore=12,
            alignment=TA_LEFT,
            fontName=base_font_bold,
            leading=19.5
        )
        
        # Nível 4: Subtítulo de subtítulo
        subsubtitle_style = ParagraphStyle(
            'SubSubtitleStyle',
            parent=styles['Heading4'],
            fontSize=12,
            textColor=colors.HexColor('#4a5568'),
            spaceAfter=4,
            spaceBefore=8,
            alignment=TA_LEFT,
            fontName=base_font_bold,
            leading=18
        )
        
        # Estilo para texto normal
        normal_style = ParagraphStyle(
            'NormalStyle',
            parent=styles['Normal'],
            fontSize=base_font_size,
            textColor=colors.HexColor('#333333'),
            spaceAfter=2,  # Reduzido para bullets mais próximos
            alignment=TA_LEFT,
            fontName=base_font,
            leading=base_leading
        )
        
        # Estilo para itálico (empresa, período)
        italic_style = ParagraphStyle(
            'ItalicStyle',
            parent=styles['Normal'],
            fontSize=base_font_size,
            textColor=colors.HexColor('#555555'),
            spaceAfter=4,
            alignment=TA_LEFT,
            fontName=base_font_italic,
            leading=base_leading
        )
        
        # Estilo para contato no cabeçalho
        contact_style = ParagraphStyle(
            'ContactStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.white,
            spaceAfter=4,
            alignment=TA_LEFT,
            fontName=base_font,
            leading=15
        )
        
        # Processa o conteúdo
        lines = cv_content.split('\n')
        
        story = []
        professional_name = None
        professional_position = None
        contact_info = []
        current_section = None
        section_content = []
        
        # Primeira passagem: extrai nome, posição e contato (até encontrar primeira seção)
        header_end_index = len(lines)
        found_first_section = False
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            if not line_stripped:
                continue
            
            # Detecta nome no formato **_[NAME]_** ou **__[NAME]__** ou **[NAME]**
            if not professional_name and line_stripped.startswith('**'):
                # Remove markdown: **_texto_** ou **__texto__** ou **texto**
                name_candidate = line_stripped
                # Remove ** do início e fim
                name_candidate = re.sub(r'^\*\*', '', name_candidate)
                name_candidate = re.sub(r'\*\*$', '', name_candidate)
                # Remove _ do início e fim (pode ter múltiplos)
                name_candidate = re.sub(r'^_+', '', name_candidate)
                name_candidate = re.sub(r'_+$', '', name_candidate)
                name_candidate = name_candidate.strip()
                
                # Verifica se não é uma seção conhecida
                is_section = name_candidate in SECTION_ICONS.keys() or any(key.lower() in name_candidate.lower() for key in SECTION_ICONS.keys())
                
                if not is_section and name_candidate and not any(icon in name_candidate for icon in ['✉', '✆', '[in]', 'http']):
                    professional_name = name_candidate
                    continue
            
            # Detecta posição no formato **__[POSITION]__** (linha após o nome)
            if professional_name and not professional_position and line_stripped.startswith('**'):
                position_candidate = line_stripped
                # Remove markdown: **__texto__** ou **texto**
                position_candidate = re.sub(r'^\*\*', '', position_candidate)
                position_candidate = re.sub(r'\*\*$', '', position_candidate)
                # Remove _ do início e fim (pode ter múltiplos)
                position_candidate = re.sub(r'^_+', '', position_candidate)
                position_candidate = re.sub(r'_+$', '', position_candidate)
                position_candidate = position_candidate.strip()
                
                # Verifica se não é uma seção conhecida
                is_section = position_candidate in SECTION_ICONS.keys() or any(key.lower() in position_candidate.lower() for key in SECTION_ICONS.keys())
                
                if not is_section and position_candidate and not any(icon in position_candidate for icon in ['✉', '✆', '[in]', 'http']):
                    professional_position = position_candidate
                    continue
            
            # Detecta primeira seção para marcar fim do cabeçalho
            if not found_first_section and line_stripped.startswith('**') and line_stripped.endswith('**') and not line_stripped.startswith('- **'):
                section_title = line_stripped.replace('**', '').strip()
                # Remove underscores se houver
                section_title = re.sub(r'^_+', '', section_title)
                section_title = re.sub(r'_+$', '', section_title)
                section_title = section_title.strip()
                # Verifica se é uma seção conhecida
                is_section = section_title in SECTION_ICONS.keys() or any(key.lower() in section_title.lower() for key in SECTION_ICONS.keys())
                
                if is_section:
                    header_end_index = i
                    found_first_section = True
                    continue
            
            # Detecta informações de contato (apenas antes da primeira seção)
            if i < header_end_index and not found_first_section:
                # Ignora linhas que são apenas separadores (****)
                if line_stripped.replace('*', '').strip() == '':
                    continue
                    
                if '✉' in line_stripped or '@' in line_stripped:
                    # Remove ✉ se presente e limpa
                    contact_line = line_stripped.replace('✉', '').strip()
                    if contact_line:
                        contact_info.append(contact_line)
                elif '✆' in line_stripped or re.search(r'\(\d{2}\)', line_stripped):
                    # Remove ✆ se presente e limpa
                    contact_line = line_stripped.replace('✆', '').strip()
                    if contact_line:
                        contact_info.append(contact_line)
                elif '[in]' in line_stripped.lower() or 'linkedin' in line_stripped.lower() or 'http' in line_stripped.lower():
                    # Limpa o formato [in] se presente
                    contact_line = re.sub(r'\[in\]\s*', '', line_stripped, flags=re.IGNORECASE).strip()
                    if contact_line:
                        contact_info.append(contact_line)
        
        # Adiciona cabeçalho colorido
        if professional_name or professional_position or contact_info:
            header_data = []
            
            # Nome (nível 1)
            if professional_name:
                name_text = escape_xml(professional_name)
                header_data.append([Paragraph(name_text, name_style)])
            
            # Posição (nível 2, menor que o nome)
            if professional_position:
                # Estilo para posição (menor que o nome)
                position_style = ParagraphStyle(
                    'PositionStyle',
                    parent=name_style,
                    fontSize=16,
                    textColor=colors.HexColor('#4a5568'),
                    spaceAfter=8,
                    spaceBefore=0,
                    alignment=TA_LEFT,
                    fontName=base_font,
                    leading=24
                )
                position_text = escape_xml(professional_position)
                header_data.append([Paragraph(position_text, position_style)])
            
            # Dados de contato
            if contact_info:
                contact_items = []
                for ci in contact_info:
                    # Processa markdown primeiro, depois escapa XML
                    processed = process_markdown(ci)
                    escaped = escape_xml(processed)
                    contact_items.append(escaped)
                contact_text = '<br/>'.join(contact_items)
                header_data.append([Paragraph(contact_text, contact_style)])
            
            if header_data:
                header_table = Table(header_data, colWidths=[doc.width])
                header_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), hex_to_color(primary_color)),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 12),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 12),
                    ('TOPPADDING', (0, 0), (-1, -1), 16),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 16),
                ]))
                story.append(header_table)
                story.append(Spacer(1, 20))
        
        # Segunda passagem: processa conteúdo (pula cabeçalho)
        i = header_end_index
        current_subtitle_group = []  # Agrupa nível 3 com seu conteúdo
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # Detecta seção principal (nível 2) - **Nome da Seção**
            if line.startswith('**') and line.endswith('**') and not line.startswith('- **'):
                section_title = line.replace('**', '').strip()
                icon = get_section_icon(section_title)
                
                # Adiciona grupo de nível 3 anterior se houver (com KeepTogether)
                if current_subtitle_group:
                    story.append(KeepTogether(current_subtitle_group))
                    current_subtitle_group = []
                
                # Adiciona conteúdo da seção anterior diretamente (sem KeepTogether - pode ser órfão)
                if current_section and section_content:
                    story.extend(section_content)
                    section_content = []
                
                # Cria parágrafo simples para seção (sem background)
                section_text = f"{icon} {section_title}"
                section_para = Paragraph(escape_xml(section_text), section_style)
                story.append(section_para)
                story.append(Spacer(1, 8))
                current_section = section_title
                i += 1
                continue
            
            # Detecta subtítulo de seção (nível 3) - - **[POSITION]**
            if line.startswith('- **') and line.endswith('**'):
                # Se havia um grupo anterior de nível 3, adiciona com KeepTogether
                if current_subtitle_group:
                    story.append(KeepTogether(current_subtitle_group))
                    current_subtitle_group = []
                
                exp_title = line.replace('- **', '').replace('**', '').strip()
                # Processa markdown primeiro, depois escapa XML
                exp_processed = process_markdown(exp_title)
                exp_text = escape_xml(exp_processed)
                # Inicia novo grupo de nível 3
                current_subtitle_group.append(Paragraph(exp_text, subtitle_style))
                current_subtitle_group.append(Spacer(1, 4))
                i += 1
                continue
            
            # Detecta subtítulo de subtítulo (nível 4) - empresa em itálico
            if line.startswith('_') and line.endswith('_'):
                company_text = line.replace('_', '').strip()
                if company_text:
                    # Processa markdown primeiro, depois escapa XML
                    company_processed = process_markdown(company_text)
                    company_escaped = escape_xml(company_processed)
                    company_para = Paragraph(company_escaped, italic_style)
                    # Adiciona ao grupo de nível 3 se existir, senão à seção
                    if current_subtitle_group:
                        current_subtitle_group.append(company_para)
                        current_subtitle_group.append(Spacer(1, 2))
                    else:
                        section_content.append(company_para)
                        section_content.append(Spacer(1, 2))
                i += 1
                continue
            
            # Detecta itens de lista com itálico (nível 4) - - *texto*
            if line.startswith('- *') and line.endswith('*'):
                italic_text = line.replace('- *', '').replace('*', '').strip()
                if italic_text:
                    # Processa markdown primeiro, depois escapa XML
                    italic_processed = process_markdown(italic_text)
                    italic_escaped = escape_xml(italic_processed)
                    italic_para = Paragraph(italic_escaped, italic_style)
                    # Adiciona ao grupo de nível 3 se existir, senão à seção
                    if current_subtitle_group:
                        current_subtitle_group.append(italic_para)
                        current_subtitle_group.append(Spacer(1, 2))
                    else:
                        section_content.append(italic_para)
                        section_content.append(Spacer(1, 2))
                i += 1
                continue
            
            # Detecta itens de lista simples (nível 4)
            if line.startswith('- '):
                list_item = line.replace('- ', '', 1).strip()
                # Processa markdown primeiro
                list_item = process_markdown(list_item)
                # Remove formatação markdown restante que não foi processada
                list_item = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', list_item)
                list_item = re.sub(r'\*(.+?)\*', r'<i>\1</i>', list_item)
                # Depois escapa XML
                list_item_escaped = escape_xml(list_item)
                bullet_text = f"• {list_item_escaped}"
                bullet_para = Paragraph(bullet_text, normal_style)
                # Adiciona ao grupo de nível 3 se existir, senão à seção
                if current_subtitle_group:
                    current_subtitle_group.append(bullet_para)
                    current_subtitle_group.append(Spacer(1, 2))  # Espaçamento reduzido
                else:
                    section_content.append(bullet_para)
                    section_content.append(Spacer(1, 2))  # Espaçamento reduzido
                i += 1
                continue
            
            # Detecta separadores
            if line.startswith('---'):
                if current_subtitle_group:
                    current_subtitle_group.append(Spacer(1, 4))
                else:
                    section_content.append(Spacer(1, 4))
                i += 1
                continue
            
            # Texto normal
            if line and not line.startswith('**') and not line.startswith('-'):
                # Processa markdown primeiro, depois escapa XML
                text_processed = process_markdown(line)
                text_escaped = escape_xml(text_processed)
                if text_escaped.strip():
                    text_para = Paragraph(text_escaped, normal_style)
                    # Adiciona ao grupo de nível 3 se existir, senão à seção
                    if current_subtitle_group:
                        current_subtitle_group.append(text_para)
                        current_subtitle_group.append(Spacer(1, 2))
                    else:
                        section_content.append(text_para)
                        section_content.append(Spacer(1, 2))
            
            i += 1
        
        # Adiciona último grupo de nível 3 se houver
        if current_subtitle_group:
            story.append(KeepTogether(current_subtitle_group))
        
        # Adiciona última seção (sem KeepTogether - pode ser órfã)
        if current_section and section_content:
            story.extend(section_content)
        
        # Gera PDF
        doc.build(story)
        
        # Retorna bytes
        buffer.seek(0)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        return pdf_bytes
        
    except ImportError:
        # Se reportlab não estiver instalado, retorna None
        st.error("Biblioteca reportlab não instalada. Execute: pip install reportlab")
        return None
    except Exception as e:
        st.error(f"Erro ao gerar PDF: {e}")
        import traceback
        st.error(traceback.format_exc())
        return None


def generate_docx_from_cv(cv_content, filename=None, primary_color="#2563eb"):
    """
    Gera um arquivo DOCX (Word) profissional a partir do conteúdo do currículo em markdown.
    
    Características:
    - 4 níveis de hierarquia tipográfica
    - Ícones para cada seção
    - Cabeçalho colorido com dados de contato
    - Cores temáticas personalizáveis
    - Espaçamento 1.5
    - Fonte tamanho 11 (normal)
    - Fonte moderna (Calibri)
    
    Args:
        cv_content: Conteúdo do currículo em markdown/texto
        filename: Nome do arquivo (opcional)
        primary_color: Cor predominante em hex (ex: "#2563eb")
    
    Returns:
        bytes: Conteúdo do DOCX em bytes
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re
        
        # Mapeamento de seções para ícones (mesmo do PDF)
        SECTION_ICONS = {
            "Resumo Profissional": "👤",
            "Resumo": "👤",
            "Experiências": "💼",
            "Experiência": "💼",
            "Formação Acadêmica": "🎓",
            "Formação": "🎓",
            "Certificações": "🏆",
            "Certificação": "🏆",
            "Cursos": "📚",
            "Cursos e Treinamentos": "📚",
            "Projetos e Consultorias Relevantes": "🚀",
            "Projetos": "🚀",
            "Hard Skills": "⚙️",
            "Habilidades Técnicas": "⚙️",
            "Soft Skills": "🤝",
            "Habilidades Comportamentais": "🤝"
        }
        
        def get_section_icon(section_name):
            """Retorna o ícone para uma seção"""
            for key, icon in SECTION_ICONS.items():
                if key.lower() in section_name.lower():
                    return icon
            return "📋"  # Ícone padrão
        
        def hex_to_rgb(hex_color):
            """Converte cor hex para RGB"""
            hex_color = hex_color.lstrip('#')
            return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        
        def process_markdown(text):
            """Processa markdown básico"""
            # Processa negrito **texto**
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            # Processa itálico *texto*
            text = re.sub(r'(?<!<b>)\*([^*<]+?)\*(?!</b>)', r'\1', text)
            # Processa links [text](url) -> text
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            return text
        
        # Cria documento Word
        doc = Document()
        
        # Configuração de fonte padrão
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Processa o conteúdo
        lines = cv_content.split('\n')
        
        professional_name = None
        professional_position = None
        contact_info = []
        header_end_index = len(lines)
        found_first_section = False
        
        # Primeira passagem: extrai nome, posição e contato
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            if not line_stripped:
                continue
            
            # Detecta nome
            if not professional_name and line_stripped.startswith('**'):
                name_candidate = line_stripped
                name_candidate = re.sub(r'^\*\*', '', name_candidate)
                name_candidate = re.sub(r'\*\*$', '', name_candidate)
                name_candidate = re.sub(r'^_+', '', name_candidate)
                name_candidate = re.sub(r'_+$', '', name_candidate)
                name_candidate = name_candidate.strip()
                
                is_section = name_candidate in SECTION_ICONS.keys() or any(key.lower() in name_candidate.lower() for key in SECTION_ICONS.keys())
                
                if not is_section and name_candidate and not any(icon in name_candidate for icon in ['✉', '✆', '[in]', 'http']):
                    professional_name = name_candidate
                    continue
            
            # Detecta posição
            if professional_name and not professional_position and line_stripped.startswith('**'):
                position_candidate = line_stripped
                position_candidate = re.sub(r'^\*\*', '', position_candidate)
                position_candidate = re.sub(r'\*\*$', '', position_candidate)
                position_candidate = re.sub(r'^_+', '', position_candidate)
                position_candidate = re.sub(r'_+$', '', position_candidate)
                position_candidate = position_candidate.strip()
                
                is_section = position_candidate in SECTION_ICONS.keys() or any(key.lower() in position_candidate.lower() for key in SECTION_ICONS.keys())
                
                if not is_section and position_candidate and not any(icon in position_candidate for icon in ['✉', '✆', '[in]', 'http']):
                    professional_position = position_candidate
                    continue
            
            # Detecta primeira seção
            if not found_first_section and line_stripped.startswith('**') and line_stripped.endswith('**') and not line_stripped.startswith('- **'):
                section_title = line_stripped.replace('**', '').strip()
                section_title = re.sub(r'^_+', '', section_title)
                section_title = re.sub(r'_+$', '', section_title)
                section_title = section_title.strip()
                is_section = section_title in SECTION_ICONS.keys() or any(key.lower() in section_title.lower() for key in SECTION_ICONS.keys())
                
                if is_section:
                    header_end_index = i
                    found_first_section = True
                    continue
            
            # Detecta contato
            if i < header_end_index and not found_first_section:
                if line_stripped.replace('*', '').strip() == '':
                    continue
                    
                if '✉' in line_stripped or '@' in line_stripped:
                    contact_line = line_stripped.replace('✉', '').strip()
                    if contact_line:
                        contact_info.append(contact_line)
                elif '✆' in line_stripped or re.search(r'\(\d{2}\)', line_stripped):
                    contact_line = line_stripped.replace('✆', '').strip()
                    if contact_line:
                        contact_info.append(contact_line)
                elif '[in]' in line_stripped.lower() or 'linkedin' in line_stripped.lower() or 'http' in line_stripped.lower():
                    contact_line = re.sub(r'\[in\]\s*', '', line_stripped, flags=re.IGNORECASE).strip()
                    if contact_line:
                        contact_info.append(contact_line)
        
        # Adiciona cabeçalho colorido
        if professional_name or professional_position or contact_info:
            # Nome
            if professional_name:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(professional_name)
                run.font.name = 'Calibri'
                run.font.size = Pt(24)
                run.font.bold = True
                run.font.color.rgb = RGBColor(26, 26, 26)
                p.space_after = Pt(6)
            
            # Posição
            if professional_position:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(professional_position)
                run.font.name = 'Calibri'
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(*hex_to_rgb('#4a5568'))
                p.space_after = Pt(8)
            
            # Contato
            if contact_info:
                for ci in contact_info:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(process_markdown(ci))
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    p.space_after = Pt(4)
            
            # Adiciona espaçamento após cabeçalho
            doc.add_paragraph().space_after = Pt(20)
        
        # Segunda passagem: processa conteúdo
        i = header_end_index
        current_section = None
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # Detecta seção principal (nível 2)
            if line.startswith('**') and line.endswith('**') and not line.startswith('- **'):
                section_title = line.replace('**', '').strip()
                section_title = re.sub(r'^_+', '', section_title)
                section_title = re.sub(r'_+$', '', section_title)
                section_title = section_title.strip()
                icon = get_section_icon(section_title)
                
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(f"{icon} {section_title}")
                run.font.name = 'Calibri'
                run.font.size = Pt(14)
                run.font.bold = True
                rgb = hex_to_rgb(primary_color)
                run.font.color.rgb = RGBColor(*rgb)
                p.space_before = Pt(16)
                p.space_after = Pt(8)
                
                current_section = section_title
                i += 1
                continue
            
            # Detecta subtítulo de seção (nível 3)
            if line.startswith('- **') and line.endswith('**'):
                exp_title = line.replace('- **', '').replace('**', '').strip()
                exp_title = process_markdown(exp_title)
                
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(exp_title)
                run.font.name = 'Calibri'
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(44, 62, 80)
                p.space_before = Pt(12)
                p.space_after = Pt(4)
                i += 1
                continue
            
            # Detecta empresa em itálico (nível 4)
            if line.startswith('_') and line.endswith('_'):
                company_text = line.replace('_', '').strip()
                if company_text:
                    company_text = process_markdown(company_text)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(company_text)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(85, 85, 85)
                    p.space_after = Pt(2)
                i += 1
                continue
            
            # Detecta itens de lista com itálico
            if line.startswith('- *') and line.endswith('*'):
                italic_text = line.replace('- *', '').replace('*', '').strip()
                if italic_text:
                    italic_text = process_markdown(italic_text)
                    p = doc.add_paragraph(style='List Bullet')
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(italic_text)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(85, 85, 85)
                    p.space_after = Pt(2)
                i += 1
                continue
            
            # Detecta itens de lista simples
            if line.startswith('- '):
                list_item = line.replace('- ', '', 1).strip()
                list_item = process_markdown(list_item)
                list_item = re.sub(r'\*\*(.+?)\*\*', r'\1', list_item)
                list_item = re.sub(r'\*(.+?)\*', r'\1', list_item)
                
                p = doc.add_paragraph(style='List Bullet')
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(list_item)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(51, 51, 51)
                p.space_after = Pt(2)
                i += 1
                continue
            
            # Detecta separadores
            if line.startswith('---'):
                doc.add_paragraph().space_after = Pt(8)
                i += 1
                continue
            
            # Texto normal
            if line and not line.startswith('**') and not line.startswith('-'):
                text_processed = process_markdown(line)
                if text_processed.strip():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(text_processed)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(51, 51, 51)
                    p.space_after = Pt(2)
            
            i += 1
        
        # Salva em buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.getvalue()
        buffer.close()
        
        return docx_bytes
        
    except ImportError:
        st.error("Biblioteca python-docx não instalada. Execute: pip install python-docx")
        return None
    except Exception as e:
        st.error(f"Erro ao gerar DOCX: {e}")
        import traceback
        st.error(traceback.format_exc())
        return None